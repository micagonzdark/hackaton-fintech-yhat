#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Twips


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Sistema_Modelos_Predictivos_Hibrido_v2.docx"

TEAL = "087F8C"
DARK_TEAL = "075965"
BLUE = "2D5BFF"
INK = "172126"
MUTED = "64727A"
LIGHT = "F2F6F7"
LIGHT_BLUE = "EAF0FF"
LIGHT_AMBER = "FFF7E8"
LIGHT_RED = "FCEEEE"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def prevent_row_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    tr_pr.append(cant_split)


def ensure_child(parent, tag):
    child = parent.find(qn(tag))
    if child is None:
        child = OxmlElement(tag)
        parent.append(child)
    return child


def set_dxa_width(parent, tag, width):
    node = ensure_child(parent, tag)
    node.set(qn("w:type"), "dxa")
    node.set(qn("w:w"), str(width))


def set_table_width(table, widths):
    widths_dxa = [round(width * 1440) for width in widths]
    widths_dxa[-1] += 9360 - sum(widths_dxa)
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    set_dxa_width(tbl_pr, "w:tblW", 9360)
    indent = ensure_child(tbl_pr, "w:tblInd")
    indent.set(qn("w:type"), "dxa")
    indent.set(qn("w:w"), "120")
    layout = ensure_child(tbl_pr, "w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        grid.append(grid_col)
    for row in table.rows:
        prevent_row_split(row)
        for index, width in enumerate(widths_dxa):
            row.cells[index].width = Twips(width)
            set_dxa_width(row.cells[index]._tc.get_or_add_tcPr(), "w:tcW", width)
            row.cells[index].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(row.cells[index])


def set_run(run, size=11, bold=False, color=INK, font="Calibri", italic=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def add_page_number(paragraph):
    run = paragraph.add_run()
    fld_char_1 = OxmlElement("w:fldChar")
    fld_char_1.set(qn("w:fldCharType"), "begin")
    instr_text = OxmlElement("w:instrText")
    instr_text.set(qn("xml:space"), "preserve")
    instr_text.text = "PAGE"
    fld_char_2 = OxmlElement("w:fldChar")
    fld_char_2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_1)
    run._r.append(instr_text)
    run._r.append(fld_char_2)
    set_run(run, size=9, color=MUTED)


def add_body(doc, text, bold_prefix=None):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.line_spacing = 1.1
    if bold_prefix and text.startswith(bold_prefix):
        set_run(paragraph.add_run(bold_prefix), bold=True)
        set_run(paragraph.add_run(text[len(bold_prefix) :]))
    else:
        set_run(paragraph.add_run(text))
    return paragraph


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    set_run(paragraph.add_run(text))
    return paragraph


def add_number(doc, text):
    paragraph = doc.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing = 1.1
    set_run(paragraph.add_run(text))
    return paragraph


def add_formula(doc, text, fill=LIGHT_BLUE):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    set_run(paragraph.add_run(text), size=10, bold=True, color=DARK_TEAL, font="Consolas")
    return paragraph


def add_callout(doc, label, text, fill=LIGHT):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.18)
    paragraph.paragraph_format.right_indent = Inches(0.18)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(10)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)
    set_run(paragraph.add_run(f"{label}: "), bold=True, color=DARK_TEAL)
    set_run(paragraph.add_run(text))
    return paragraph


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table, widths)
    set_repeat_table_header(table.rows[0])
    for index, header in enumerate(headers):
        cell = table.rows[0].cells[index]
        set_cell_shading(cell, TEAL)
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_run(paragraph.add_run(header), size=9.5, bold=True, color=WHITE)
    for row_index, values in enumerate(rows):
        cells = table.add_row().cells
        for index, value in enumerate(values):
            set_cell_shading(cells[index], WHITE if row_index % 2 == 0 else LIGHT)
            paragraph = cells[index].paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            set_run(paragraph.add_run(str(value)), size=9.5)
        set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.49)
    section.footer_distance = Inches(0.49)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color, before, after in (
        ("Title", 25, INK, 0, 5),
        ("Subtitle", 13, MUTED, 0, 18),
        ("Heading 1", 17, TEAL, 16, 8),
        ("Heading 2", 13, DARK_TEAL, 12, 6),
        ("Heading 3", 11.5, BLUE, 8, 4),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.bold = style_name != "Subtitle"
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    set_run(header.add_run("RENTA ESTABLE  |  SISTEMA HÍBRIDO v2.0"), size=9, bold=True, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run(footer.add_run("Documento técnico  |  Página "), size=9, color=MUTED)
    add_page_number(footer)


def build_document():
    doc = Document()
    configure_document(doc)

    kicker = doc.add_paragraph()
    kicker.paragraph_format.space_after = Pt(6)
    set_run(kicker.add_run("DOCUMENTO TÉCNICO v2.0"), size=10, bold=True, color=TEAL)

    title = doc.add_paragraph(style="Title")
    set_run(title.add_run("Sistema híbrido de modelos predictivos y adelantos"), size=25, bold=True, color=INK)
    subtitle = doc.add_paragraph(style="Subtitle")
    set_run(
        subtitle.add_run("Financiamiento responsable para anfitriones de alquiler temporario"),
        size=13,
        color=MUTED,
    )

    metadata = [
        ("Producto", "FlowRent: adelanto opt-in de cobros futuros"),
        ("Versión", "2.0 — arquitectura híbrida"),
        ("Fecha", "6 de junio de 2026"),
        ("Estado", "MVP explicable con ruta de calibración"),
    ]
    add_table(doc, ["Campo", "Definición"], metadata, [1.35, 5.15])

    add_callout(
        doc,
        "Decisión principal",
        "No reemplazar el score experto por modelos no entrenados. El MVP conserva reglas, P10 y retención; agrega una PD proxy claramente etiquetada, pricing explicable y un plan de calibración con resultados reales.",
        LIGHT_BLUE,
    )

    doc.add_heading("Resumen ejecutivo", level=1)
    add_body(
        doc,
        "El sistema ejecuta capas complementarias. Primero excluye riesgos que no deben compensarse con un promedio. Luego resume la solidez del host con un score explicable, lo convierte temporalmente en una PD proxy, calcula cuánto flujo futuro puede recuperarse y finalmente compara el fee contractual con un precio sugerido.",
    )
    add_formula(
        doc,
        "Reglas duras → Score experto → PD → P10 + retención → Monto → Pricing → Monitoreo",
    )
    add_body(
        doc,
        "El diferencial no es solamente predecir quién pagará. La plataforma observa reservas, fechas de payout y reputación, y puede recuperar el adelanto dentro del mismo flujo de pagos.",
    )

    doc.add_page_break()
    doc.add_heading("1. Arquitectura completa", level=1)
    add_table(
        doc,
        ["Capa", "Pregunta", "Salida"],
        [
            ("0. Reglas duras", "¿Existe una condición excluyente?", "Rechazo o revisión"),
            ("1. Riesgo", "¿Qué tan probable es un incumplimiento?", "Score + PD 12m"),
            ("2. Recuperabilidad", "¿Cuánto puede retenerse sin sobreexponer?", "Tope P10 y cobertura visible"),
            ("3. Pricing", "¿Qué precio cubre costo, riesgo y margen?", "Fee sugerido"),
            ("4. Monitoreo", "¿Qué ocurrió realmente?", "Resultado, recupero y recalibración"),
        ],
        [1.25, 3.2, 2.05],
    )
    add_body(
        doc,
        "La separación evita tres errores frecuentes: usar el score como monto, llamar predictivo a un score manual y fijar un precio sin conectar PD, LGD, plazo y costo de fondeo.",
    )

    doc.add_heading("0. Reglas duras", level=2)
    for item in (
        "Fraude o suspensión de cuenta.",
        "Identidad o cuenta bancaria no verificadas.",
        "Habilitación legal pendiente cuando sea exigible.",
        "Ausencia de control contractual y operativo sobre payouts.",
        "Consentimiento, KYC o condiciones regulatorias incompletas.",
    ):
        add_bullet(doc, item)
    add_callout(
        doc,
        "Principio",
        "Una regla excluyente no puede arreglarse subiendo manualmente el score. El sistema debe conservar el rechazo y explicar la razón.",
        LIGHT_AMBER,
    )

    doc.add_heading("1. Score experto y PD", level=2)
    add_body(
        doc,
        "El scorecard por clusters sigue siendo útil en lanzamiento: ordena señales, permite discusión con riesgo y produce razones entendibles. Sus pesos son hipótesis expertas, no coeficientes estimados.",
    )
    add_table(
        doc,
        ["Cluster", "Peso", "Función"],
        [
            ("Solidez del host", "15%", "Historial, verificaciones y conducta"),
            ("Calidad de propiedad", "15%", "Capacidad futura del activo"),
            ("Historial de ingresos", "25%", "Volumen, estabilidad y estacionalidad"),
            ("Reputación y operación", "15%", "Cancelaciones, disputas y respuesta"),
            ("Reservas futuras", "20%", "Flujo ya observable"),
            ("Riesgo de mercado", "10%", "Clima, demanda y regulación"),
        ],
        [2.0, 0.75, 3.75],
    )

    doc.add_page_break()
    doc.add_heading("2. Modelo 1 — Probabilidad de default", level=1)
    doc.add_heading("Etapa MVP: PD proxy", level=2)
    add_body(
        doc,
        "Sin operaciones maduras y defaults suficientes, no existe una regresión logística entrenable con validez. Para conectar el score con monto y economía, la demo utiliza una transformación monotónica explícitamente denominada PD proxy.",
    )
    add_formula(
        doc,
        "PD_proxy_pct = clamp(50 × exp(-0,06 × (score - 40)), 1,5%, 45%)",
    )
    add_callout(
        doc,
        "Advertencia",
        "La PD proxy no es una probabilidad calibrada. Sirve para demostrar arquitectura, sensibilidad y unit economics; no debe presentarse como performance histórica.",
        LIGHT_RED,
    )

    doc.add_heading("Etapa calibrada: regresión logística", level=2)
    add_body(
        doc,
        "Cuando exista una cantidad suficiente de operaciones resueltas y eventos de default, la proxy se reemplaza por regresión logística. Las variables monetarias deben ajustarse por inflación y transformarse o normalizarse para evitar que la escala domine el resultado.",
    )
    add_formula(doc, "PD = 1 / (1 + exp(-(β₀ + β₁x₁ + ... + βₙxₙ)))")
    add_body(doc, "Definición propuesta de la etiqueta:")
    add_formula(
        doc,
        "default_120d = 1 si no se recuperó ≥95% del total contractual 120 días después de la fecha estimada de finalización",
        LIGHT_AMBER,
    )
    for item in (
        "Los casos todavía no maduros se excluyen del entrenamiento; no se etiquetan como pagados.",
        "La aprobación o rechazo anterior no puede usarse como etiqueta, porque sólo replicaría decisiones pasadas.",
        "Se documentan coeficientes, variables, datos faltantes, overrides y versión de modelo.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Validación de PD", level=2)
    add_table(
        doc,
        ["Dimensión", "Métrica mínima", "Pregunta"],
        [
            ("Calibración", "Curva por bandas, Brier", "¿Un 8% observado se parece a 8%?"),
            ("Discriminación", "AUC/Gini o KS", "¿Ordena mejor y peor riesgo?"),
            ("Estabilidad", "PSI y drift", "¿Cambió la población o la señal?"),
            ("Negocio", "Pérdida real vs esperada", "¿La PD sirve para decidir y provisionar?"),
        ],
        [1.35, 2.1, 3.05],
    )

    doc.add_page_break()
    doc.add_heading("3. Modelo 2 — Monto y recuperabilidad", level=1)
    add_body(
        doc,
        "El monto no debe aprenderse únicamente de operaciones pagadas ni usar como etiqueta el monto repagado. Eso introduce sesgo de supervivencia y reproduce límites anteriores. La arquitectura estima ingresos futuros conservadores y aplica una política de recuperabilidad.",
    )
    doc.add_heading("Ingreso futuro conservador", level=2)
    add_formula(doc, "P10_ajustado = P10_ingreso_neto × (1 - caída_temporada)")
    add_body(
        doc,
        "En el MVP el P10 está precargado. Con datos suficientes se estima mediante regresión cuantílica o series temporales y se valida midiendo la frecuencia real de incumplimiento del percentil.",
    )

    doc.add_heading("Fórmula híbrida de monto", level=2)
    add_formula(doc, "tope_recuperable = P10_ajustado × retención / (1 + fee_contractual)")
    add_formula(doc, "tope_ajustado_PD = tope_recuperable × (1 - PD)")
    add_formula(
        doc,
        "monto_recomendado = min(monto_solicitado, tope_ajustado_PD, límite_política)",
    )
    add_body(
        doc,
        "Las reservas confirmadas se calculan por separado como cobertura visible. Esta métrica muestra cuánto del total contractual ya puede observarse, sin confundir reservas actuales con todo el flujo futuro esperado.",
    )

    doc.add_heading("Política de decisión MVP", level=2)
    add_table(
        doc,
        ["Condición", "Decisión"],
        [
            ("Regla dura, score < 50 o PD ≥ 35%", "Rechazado"),
            ("Score < 65 o PD ≥ 15%", "Línea piloto"),
            ("Monto recomendado ≥ 95% del pedido", "Aprobado"),
            ("Monto recomendado ≥ 45% del pedido", "Aprobación parcial"),
            ("Monto positivo restante", "Línea piloto"),
        ],
        [4.25, 2.25],
    )

    doc.add_page_break()
    doc.add_heading("4. Modelo 3 — Pricing explicable", level=1)
    add_body(
        doc,
        "El producto conserva un fee contractual visible para el host y calcula, en paralelo, un fee sugerido. Esta comparación permite detectar operaciones que no cubren fondeo y riesgo sin cambiar automáticamente el contrato.",
    )
    add_formula(doc, "prima_riesgo_anual = (PD × LGD) / (1 - PD)")
    add_formula(
        doc,
        "tasa_anual_sugerida = fondeo_anual + margen_objetivo_anual + prima_riesgo_anual",
    )
    add_formula(
        doc,
        "fee_sugerido = tasa_anual_sugerida × meses / 12 + costo_operativo_por_operación",
    )
    add_callout(
        doc,
        "Distinción clave",
        "La tasa es anual; el fee es el costo total de una operación de plazo concreto. No deben compararse como si fueran la misma unidad.",
        LIGHT_BLUE,
    )

    doc.add_heading("Unit economics individual", level=2)
    add_formula(doc, "pérdida_esperada_host = principal × PD_host × LGD")
    add_formula(
        doc,
        "contribución = fee - costo_fondeo_por_plazo - pérdida_esperada - costo_operativo",
    )
    add_body(
        doc,
        "La cartera deja de usar una única PD global. Cada host aporta su propia pérdida esperada y la sensibilidad aplica stress proporcional a todas las PD individuales.",
    )

    doc.add_heading("Ejemplo: María, Bariloche", level=2)
    add_table(
        doc,
        ["Variable", "Valor base"],
        [
            ("Score", "89,0"),
            ("PD proxy 12m", "2,64%"),
            ("P10 futuro", "ARS 11.800.000"),
            ("Retención / fee contractual", "28% / 10%"),
            ("Tope recuperable", "ARS 3.003.636"),
            ("Tope ajustado por PD", "ARS 2.924.242"),
            ("Monto pedido / recomendado", "ARS 2.800.000 / ARS 2.800.000"),
            ("Fee sugerido", "12,49%"),
        ],
        [2.65, 3.85],
    )

    doc.add_page_break()
    doc.add_heading("5. Datos, monitoreo y gobernanza", level=1)
    add_body(
        doc,
        "La mejora del sistema depende de capturar decisiones y resultados con definiciones estables. Guardar sólo la oferta final no alcanza.",
    )
    add_table(
        doc,
        ["Entidad", "Campos críticos"],
        [
            ("Snapshot de variables", "Valores usados, fecha, fuente y datos faltantes"),
            ("Decisión", "Score, PD, reglas activadas, thresholds y override"),
            ("Oferta", "Principal, fee, retención, plazo y total contractual"),
            ("Resultado", "Payouts retenidos, cancelaciones, recupero y default"),
            ("Modelo", "Versión, método, fecha de entrenamiento y validación"),
        ],
        [1.65, 4.85],
    )
    add_body(doc, "Controles mínimos:")
    for item in (
        "Versionar scorecard, PD, P10, thresholds y fórmulas de pricing.",
        "Registrar razones de decisión y toda revisión manual.",
        "Monitorear drift, concentración por mercado y performance por cohorte.",
        "Permitir corrección de datos y revisión de decisiones adversas.",
        "Separar marketplace, pagos y entidad financiadora cuando corresponda.",
    ):
        add_bullet(doc, item)

    doc.add_heading("Roadmap por madurez de datos", level=2)
    add_table(
        doc,
        ["Etapa", "Condición de entrada", "Modelo"],
        [
            ("MVP", "Sin eventos suficientes", "Scorecard + PD proxy + P10 precargado"),
            ("Calibración", "Operaciones maduras y defaults suficientes", "Logística calibrada y backtesting"),
            ("Forecast", "Series históricas suficientes", "P10/P50/P90 por cuantiles"),
            ("Optimización", "Volumen y variación de ofertas", "Límites y pricing optimizados"),
        ],
        [1.1, 2.75, 2.65],
    )
    add_callout(
        doc,
        "Criterio de avance",
        "Los meses transcurridos no garantizan validez. Cada etapa comienza cuando existen suficientes observaciones maduras, eventos y estabilidad para validarla fuera de muestra.",
        LIGHT_AMBER,
    )

    doc.add_page_break()
    doc.add_heading("6. Implementación actual en la demo", level=1)
    add_body(
        doc,
        "La versión híbrida ya está reflejada en el simulador y en el esquema de datos del repositorio.",
    )
    add_table(
        doc,
        ["Componente", "Implementación actual"],
        [
            ("UI", "Muestra PD proxy, fee sugerido, tope P10 y tope ajustado por PD"),
            ("Monto", "Aplica P10, retención, fee contractual y buffer por PD"),
            ("Reglas", "Mantiene rechazo aunque se modifiquen scores manualmente"),
            ("Cartera", "Calcula pérdida esperada con PD individual ponderada"),
            ("Sensibilidad", "Stress de fondeo, PD individual y LGD"),
            ("Datos", "Persiste método de PD, topes y fee sugerido"),
        ],
        [1.45, 5.05],
    )

    doc.add_heading("Mensajes recomendados para pitch", level=2)
    for text in (
        "No adelantamos contra optimismo: usamos P10, retención y stress.",
        "El score explica; la PD cuantifica riesgo; el flujo futuro define el monto.",
        "La PD actual es una proxy experta y se reemplaza por una PD calibrada con resultados reales.",
        "El fee se contrasta con fondeo, plazo, PD y LGD para mostrar cuándo el negocio funciona.",
        "Para producción, montos mayores o flujo no confirmado requieren partner regulado o encuadre validado.",
    ):
        add_bullet(doc, text)

    doc.add_heading("Conclusión", level=1)
    add_body(
        doc,
        "La arquitectura híbrida es más sólida que cualquiera de los dos enfoques por separado. Conserva la recuperabilidad, transparencia y prudencia del adelanto de payouts; incorpora PD individual, pricing y aprendizaje progresivo; y evita presentar supuestos expertos como modelos estadísticos ya validados.",
    )
    add_callout(
        doc,
        "Frase final",
        "FlowRent convierte datos operativos en adelantos responsables: primero limita lo que puede recuperarse, luego ajusta por riesgo y finalmente muestra si el precio cubre la operación.",
        LIGHT_BLUE,
    )

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_document()
